"""Minimal Markdown -> .docx converter for the project reports.

Supports the subset used in the reports: #/##/### headings, | tables |,
- bullets, 1. numbers, ``` fenced code/diagram blocks ```, > blockquotes,
paragraphs, and inline **bold** / `code` / [text](url).

Usage: python md2docx.py <input.md> <output.docx>
"""
from __future__ import annotations

import re
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

_HEADER_FILL = "1F4E79"
_HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)
_CODE_FONT = "Consolas"
_BODY_EAST_ASIAN = "Microsoft YaHei"


def _set_cell_shading(cell, fill_hex: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): fill_hex})
    tcpr.append(shd)


def _add_inline(paragraph, text: str, *, bold: bool = False, color=None) -> None:
    """Add runs, parsing **bold**, `code`, and [label](url) -> label."""
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)  # links -> label
    # split on bold and inline code while keeping delimiters
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = _CODE_FONT
        else:
            run = paragraph.add_run(part)
            run.bold = bold
        if color is not None:
            run.font.color.rgb = color
        _set_run_eastasia(run)


def _set_run_eastasia(run) -> None:
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), _BODY_EAST_ASIAN)


def _emit_table(doc, rows: list[list[str]]) -> None:
    if not rows:
        return
    ncol = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=ncol)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, raw in enumerate(rows):
        cells = table.add_row().cells
        for ci in range(ncol):
            text = raw[ci] if ci < len(raw) else ""
            para = cells[ci].paragraphs[0]
            if ri == 0:
                _set_cell_shading(cells[ci], _HEADER_FILL)
                _add_inline(para, text, bold=True, color=_HEADER_TEXT)
            else:
                _add_inline(para, text)
            for run in para.runs:
                run.font.size = Pt(9)


def _emit_code(doc, lines: list[str]) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(8)
    para.paragraph_format.left_indent = Pt(6)
    for i, ln in enumerate(lines):
        if i:
            para.add_run().add_break()
        run = para.add_run(ln if ln else " ")
        run.font.name = _CODE_FONT
        run.font.size = Pt(8)
        _set_run_eastasia(run)


def convert(md_path: str, docx_path: str) -> None:
    lines = open(md_path, encoding="utf-8").read().splitlines()
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        # fenced code / diagram block
        if line.lstrip().startswith("```"):
            block: list[str] = []
            i += 1
            while i < n and not lines[i].lstrip().startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            _emit_code(doc, block)
            continue
        # table
        if line.strip().startswith("|") and i + 1 < n and re.match(r"^\s*\|[-:\s|]+\|\s*$", lines[i + 1]):
            rows: list[list[str]] = []
            header = [c.strip() for c in line.strip().strip("|").split("|")]
            rows.append(header)
            i += 2  # skip header + separator
            while i < n and lines[i].strip().startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            _emit_table(doc, rows)
            continue
        # headings
        m = re.match(r"^(#{1,4})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            doc.add_heading(re.sub(r"[*`]", "", m.group(2)).strip(), level=min(level, 4))
            i += 1
            continue
        # blockquote
        if line.startswith(">"):
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Pt(12)
            _add_inline(para, line.lstrip("> ").rstrip())
            for run in para.runs:
                run.italic = True
                run.font.size = Pt(10)
            i += 1
            continue
        # horizontal rule
        if re.match(r"^\s*---+\s*$", line):
            i += 1
            continue
        # bullet list
        m = re.match(r"^\s*[-*]\s+(.*)$", line)
        if m:
            para = doc.add_paragraph(style="List Bullet")
            _add_inline(para, m.group(1))
            i += 1
            continue
        # numbered list
        m = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if m:
            para = doc.add_paragraph(style="List Number")
            _add_inline(para, m.group(1))
            i += 1
            continue
        # blank
        if not line.strip():
            i += 1
            continue
        # paragraph
        para = doc.add_paragraph()
        _add_inline(para, line)
        i += 1

    doc.save(docx_path)
    print(f"wrote {docx_path}")


if __name__ == "__main__":
    convert(sys.argv[1], sys.argv[2])
